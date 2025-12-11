# resync/core/file_ingestor.py

import os
import re
import shutil
from collections.abc import Iterator
from pathlib import Path

import docx
import openpyxl
import pypdf  # Corrected import for pypdf
from openpyxl.utils.exceptions import InvalidFileException

from resync.core.exceptions import FileProcessingError, KnowledgeGraphError
from resync.core.interfaces import IFileIngestor, IKnowledgeGraph
from resync.core.structured_logger import get_logger
from resync.settings import settings

logger = get_logger(__name__)


def is_path_protected(file_path: Path) -> bool:
    """
    Check if a file path is within a protected directory.

    Args:
        file_path: Path to check for protection

    Returns:
        True if the path is within a protected directory, False otherwise
    """
    if not file_path.exists():
        return False

    file_path = file_path.resolve()

    for protected_dir in settings.PROTECTED_DIRECTORIES:
        protected_dir = Path(protected_dir).resolve()
        try:
            # Check if the file path is within the protected directory
            file_path.relative_to(protected_dir)
            return True
        except ValueError:
            # Path is not within the protected directory
            continue

    return False


def is_path_in_knowledge_base(file_path: Path) -> bool:
    """
    Check if a file path is within a knowledge base directory.

    Args:
        file_path: Path to check for knowledge base inclusion

    Returns:
        True if the path is within a knowledge base directory, False otherwise
    """
    if not file_path.exists():
        return False

    file_path = file_path.resolve()

    for knowledge_dir in settings.KNOWLEDGE_BASE_DIRS:
        knowledge_dir = Path(knowledge_dir).resolve()
        try:
            # Check if the file path is within the knowledge base directory
            file_path.relative_to(knowledge_dir)
            return True
        except ValueError:
            # Path is not within the knowledge base directory
            continue

    return False


# --- Text Chunking --- #


def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> Iterator[str]:
    """Splits a long text into smaller chunks with overlap."""
    if not text:
        return

    start = 0
    while start < len(text):
        end = start + chunk_size
        yield text[start:end]
        start += chunk_size - chunk_overlap


# --- File Readers --- #


def read_pdf(file_path: Path) -> str:
    """Extracts text from a PDF file."""
    logger.info("reading_pdf_file", file_path=str(file_path))
    try:
        reader = pypdf.PdfReader(file_path)
        return "".join(page.extract_text() for page in reader.pages if page.extract_text())
    except FileNotFoundError as e:
        logger.error("pdf_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error("permission_denied_reading_pdf", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except pypdf.errors.PdfReadError as e:
        logger.error("pdf_read_error", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except ValueError as e:
        logger.error("invalid_pdf_content", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch any other pypdf or system errors
        logger.critical(
            "unexpected_error_reading_pdf",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process PDF {file_path}") from e


def read_json(file_path: Path) -> str:
    """Extracts text from a JSON file."""
    logger.info("reading_json_file", file_path=str(file_path))
    try:
        import json

        with open(file_path, encoding="utf-8") as f:
            data = json.load(f)

        # Convert JSON to readable text format with optimized string operations
        if isinstance(data, dict):
            # Use list comprehension and join for better performance
            text_parts = [
                f"{key}: {json.dumps(value, indent=2) if isinstance(value, (dict, list)) else value}"
                for key, value in data.items()
            ]
            return "\n".join(text_parts)
        if isinstance(data, list):
            return json.dumps(data, indent=2)
        return str(data)

    except FileNotFoundError as e:
        logger.error("json_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error("permission_denied_reading_json", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except json.JSONDecodeError as e:
        logger.error("invalid_json_format", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except UnicodeDecodeError as e:
        logger.error("encoding_error_reading_json", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch other potential OS or parsing errors
        logger.critical(
            "unexpected_error_reading_json",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process JSON {file_path}") from e


def read_txt(file_path: Path) -> str:
    """Extracts text from a plain text file."""
    logger.info("reading_text_file", file_path=str(file_path))
    try:
        return file_path.read_text(encoding="utf-8")
    except FileNotFoundError as e:
        logger.error("text_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_text_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except UnicodeDecodeError as e:
        logger.error("encoding_error_reading_text_file", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except OSError as e:
        logger.critical(
            "unexpected_os_error_reading_text_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process text file {file_path}") from e


def read_doc(file_path: Path) -> str:
    """Extracts text from a DOC file (older Word format)."""
    logger.info("reading_doc_file", file_path=str(file_path))
    try:
        # Try to use python-docx first (it can sometimes handle .doc files)
        try:
            doc = docx.Document(file_path)
            text = "\n".join(para.text for para in doc.paragraphs if para.text)
            if text.strip():
                return text
        except Exception as e:  # docx library might fail on .doc
            # Log that docx failed, will try fallback method
            logger.debug(f"docx library failed on DOC file, trying fallback: {e}")

        # Fallback: try to use antiword or similar tool if available
        # For now, return a message indicating the file type is not fully supported
        logger.warning("doc_file_requires_manual_processing", file_path=str(file_path))
        return f"[DOC file: {file_path.name} - Manual processing may be required]"

    except FileNotFoundError as e:
        logger.error("doc_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error("permission_denied_reading_doc_file", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "unexpected_error_reading_doc_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process DOC file {file_path}") from e


def read_xls(file_path: Path) -> str:
    """Extracts text from an XLS file (older Excel format)."""
    logger.info("reading_xls_file", file_path=str(file_path))
    try:
        # Try to use openpyxl first (it can sometimes handle .xls files)
        try:
            workbook = openpyxl.load_workbook(file_path, read_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    # Convert row to text, filtering out None values
                    row_text = " | ".join(str(cell) for cell in row if cell is not None)
                    if row_text.strip():
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line between sheets

            workbook.close()
            return "\n".join(text_parts)

        except Exception as e:  # openpyxl might fail on .xls
            # Log that openpyxl failed, will try xlrd fallback
            logger.debug(f"openpyxl library failed on XLS file, trying xlrd: {e}")

        # Fallback: try to use xlrd if available, or return message
        try:
            import xlrd

            workbook = xlrd.open_workbook(file_path)
            text_parts = []

            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                text_parts.append(f"Sheet: {sheet_name}")

                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    # Convert row to text, filtering out None values
                    row_text = " | ".join(
                        str(cell) for cell in row if cell is not None and str(cell).strip()
                    )
                    if row_text.strip():
                        text_parts.append(row_text)

                text_parts.append("")  # Empty line between sheets

            return "\n".join(text_parts)

        except ImportError:
            logger.warning("xls_file_requires_xlrd_library", file_path=str(file_path))
            return f"[XLS file: {file_path.name} - Install xlrd for better support]"
        except Exception as e:
            logger.error("error_processing_xls_file", file_path=str(file_path), error=str(e), exc_info=True)
            return f"[XLS file: {file_path.name} - Processing error: {e}]"

    except FileNotFoundError as e:
        logger.error("xls_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error("permission_denied_reading_xls_file", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "unexpected_error_reading_xls_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process XLS file {file_path}") from e


def read_md(file_path: Path) -> str:
    """Extracts text from a Markdown file."""
    logger.info("reading_markdown_file", file_path=str(file_path))
    try:
        return file_path.read_text(encoding="utf-8")
    except FileNotFoundError as e:
        logger.error("markdown_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except UnicodeDecodeError as e:
        logger.error(
            "encoding_error_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except OSError as e:
        logger.critical(
            "unexpected_os_error_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process Markdown file {file_path}") from e


def read_html(file_path: Path) -> str:
    """
    Extracts text from an HTML file using BeautifulSoup4.

    Optimized for IBM/HCL TWS documentation pages.
    Removes navigation, headers, footers, and preserves code blocks.
    """
    logger.info("reading_html_file", file_path=str(file_path))
    try:
        # Try to use BeautifulSoup4
        try:
            from bs4 import BeautifulSoup

            html_content = file_path.read_text(encoding="utf-8")
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for element in soup(["script", "style", "noscript"]):
                element.decompose()

            # Remove navigation elements
            nav_selectors = [
                "nav",
                "header",
                "footer",
                ".navigation",
                ".nav",
                ".navbar",
                ".sidebar",
                ".side-nav",
                ".sidenav",
                ".toc",
                ".table-of-contents",
                ".breadcrumb",
                ".breadcrumbs",
                "#header",
                "#footer",
                "#nav",
                "#sidebar",
                "[role='navigation']",
                "[role='banner']",
                ".ibm-masthead",
                ".ibm-footer",
                ".hcl-header",
                ".hcl-footer",
            ]

            for selector in nav_selectors:
                for element in soup.select(selector):
                    element.decompose()

            # Try to find main content
            main_content = None
            main_selectors = [
                "main",
                "article",
                ".content",
                ".main-content",
                ".doc-content",
                "#content",
                "#main-content",
                "#article",
                ".ibm-content",
                ".hcl-content",
                "[role='main']",
            ]

            for selector in main_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break

            # Use main content or fallback to body
            target = main_content or soup.find("body") or soup

            # Preserve code blocks
            for code in target.find_all(["code", "pre"]):
                code_text = code.get_text()
                code.replace_with(f"\n```\n{code_text}\n```\n")

            # Get text with separator
            text = target.get_text(separator="\n")

            # Clean up excessive whitespace
            text = re.sub(r"\n{3,}", "\n\n", text)
            text = re.sub(r"[ \t]+", " ", text)

            return text.strip()

        except ImportError:
            # Fallback: basic HTML parsing without BeautifulSoup
            logger.warning(
                "beautifulsoup4_not_installed",
                file_path=str(file_path),
                message="Install beautifulsoup4 for better HTML parsing",
            )

            html_content = file_path.read_text(encoding="utf-8")

            # Remove script and style tags
            html_content = re.sub(
                r"<(script|style)[^>]*>.*?</\1>", "", html_content, flags=re.DOTALL | re.IGNORECASE
            )

            # Remove HTML tags
            text = re.sub(r"<[^>]+>", " ", html_content)

            # Decode HTML entities
            import html

            text = html.unescape(text)

            # Clean up whitespace
            text = re.sub(r"\s+", " ", text)

            return text.strip()

    except FileNotFoundError as e:
        logger.error("html_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_html_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except UnicodeDecodeError as e:
        # Try with different encodings
        for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
            try:
                text = file_path.read_text(encoding=encoding)
                logger.info(
                    "html_file_read_with_fallback_encoding",
                    file_path=str(file_path),
                    encoding=encoding,
                )
                return text
            except UnicodeDecodeError:
                continue
        logger.error(
            "encoding_error_reading_html_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except Exception as e:
        logger.critical(
            "unexpected_error_reading_html_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process HTML file {file_path}") from e
    logger.info("reading_markdown_file", file_path=str(file_path))
    try:
        return file_path.read_text(encoding="utf-8")
    except FileNotFoundError as e:
        logger.error("markdown_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except UnicodeDecodeError as e:
        logger.error(
            "encoding_error_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except OSError as e:
        logger.critical(
            "unexpected_os_error_reading_markdown_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process Markdown file {file_path}") from e


def read_docx(file_path: Path) -> str:
    """Extracts text from a DOCX file."""
    logger.info("reading_docx_file", file_path=str(file_path))
    try:
        doc = docx.Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs if para.text)
    except FileNotFoundError as e:
        logger.error("docx_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_docx_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""  # Corrected import for PackageNotFoundError
    except docx.opc.exceptions.PackageNotFoundError as e:
        logger.error("invalid_docx_package", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except ValueError as e:
        logger.error("invalid_docx_content", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "unexpected_error_reading_docx_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process DOCX file {file_path}") from e


def read_excel(file_path: Path) -> str:
    """Extracts text from an XLSX file, iterating through all sheets and cells."""
    logger.info("reading_excel_file", file_path=str(file_path))
    text_parts = []
    try:
        workbook = openpyxl.load_workbook(file_path, read_only=True)
        for sheetname in workbook.sheetnames:
            sheet = workbook[sheetname]
            text_parts.append(f"Sheet: {sheetname}\n")
            for row in sheet.iter_rows():
                row_texts = []
                for cell in row:
                    if cell.value is not None:
                        row_texts.append(str(cell.value))
                if row_texts:
                    text_parts.append(" | ".join(row_texts))
        return "\n".join(text_parts)
    except FileNotFoundError as e:
        logger.error("excel_file_not_found", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except PermissionError as e:
        logger.error(
            "permission_denied_reading_excel_file",
            file_path=str(file_path),
            error=str(e),
        )
        return ""
    except InvalidFileException as e:
        logger.error("invalid_excel_file", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except ValueError as e:
        logger.error("invalid_excel_content", file_path=str(file_path), error=str(e), exc_info=True)
        return ""
    except Exception as e:  # Catch other potential library or system errors
        logger.critical(
            "unexpected_error_reading_excel_file",
            file_path=str(file_path),
            error=str(e),
            exc_info=True,
        )
        raise FileProcessingError(f"Failed to process Excel file {file_path}") from e


# --- Main Ingestion Logic --- #


class FileIngestor(IFileIngestor):
    """
    Service for ingesting files into the knowledge graph.

    This class handles file uploads, saving, and processing for RAG.
    """

    def __init__(self, knowledge_graph: IKnowledgeGraph):
        """
        Initialize the FileIngestor with dependencies.

        Args:
            knowledge_graph: The knowledge graph service to store extracted content
        """
        self.knowledge_graph = knowledge_graph
        self.rag_directory = settings.BASE_DIR / "rag"
        self.file_readers = {
            ".pdf": read_pdf,
            ".docx": read_docx,
            ".xlsx": read_excel,
            ".md": read_md,
            ".json": read_json,
            ".txt": read_txt,
            ".doc": read_doc,
            ".xls": read_xls,
            ".html": read_html,
            ".htm": read_html,
        }
        # Ensure the RAG directory exists
        self.rag_directory.mkdir(exist_ok=True)
        logger.info("file_ingestor_initialized", rag_directory=str(self.rag_directory))

    async def save_uploaded_file(self, file_name: str, file_content) -> Path:
        """
        Saves an uploaded file to the RAG directory with proper sanitization.

        Args:
            file_name: The original filename
            file_content: A file-like object containing the content

        Returns:
            Path to the saved file

        Raises:
            FileProcessingError: If the file cannot be saved
        """
        # Sanitize the filename to prevent security risks like path traversal
        basename = os.path.basename(file_name)
        # Remove any characters that aren't alphanumeric, dots, underscores, or hyphens
        safe_filename = re.sub(r"[^\w\-_.]", "", basename)
        if not safe_filename:
            raise FileProcessingError("Invalid filename.")

        destination = self.rag_directory / safe_filename

        try:
            logger.info("saving_uploaded_file", filename=safe_filename)
            with destination.open("wb") as buffer:
                shutil.copyfileobj(file_content, buffer)
            logger.info("successfully_saved_file", destination=str(destination))
            return destination
        except OSError as e:
            logger.critical(
                "failed_to_save_uploaded_file",
                filename=safe_filename,
                error=str(e),
                exc_info=True,
            )
            raise FileProcessingError(f"Could not save file due to OS error: {e}") from e

    async def ingest_file(self, file_path: Path) -> bool:
        """
        Ingests a single file into the knowledge graph.

        Args:
            file_path: Path to the file to ingest

        Returns:
            True if ingestion was successful, False otherwise
        """
        if not file_path.exists():
            logger.warning("file_not_found_for_ingestion", file_path=str(file_path))
            return False

        # Check if file is in knowledge base directories
        if not is_path_in_knowledge_base(file_path):
            logger.warning("file_not_in_knowledge_base_directories", file_path=str(file_path))
            return False

        # Check if file is in protected directories (should not be deleted during processing)
        if is_path_protected(file_path):
            logger.info("file_in_protected_directory", file_path=str(file_path))

        file_ext = file_path.suffix.lower()
        reader = self.file_readers.get(file_ext)

        if not reader:
            logger.warning("unsupported_file_type", file_extension=file_ext)
            return False

        # Read the file content
        content = reader(file_path)
        if not content:
            logger.warning("no_content_extracted", file_path=str(file_path))
            return False

        # Chunk the content and add to knowledge graph
        chunks = list(chunk_text(content))
        chunk_count = 0
        for i, chunk in enumerate(chunks):
            try:
                metadata = {
                    "source_file": str(file_path.name),
                    "chunk_index": i + 1,
                    "total_chunks": len(chunks),
                }
                # Here we add the chunk to the knowledge graph
                await self.knowledge_graph.add_content(content=chunk, metadata=metadata)
                chunk_count += 1
            except KnowledgeGraphError as e:
                logger.error(
                    "knowledge_graph_error_adding_chunk",
                    chunk_index=i + 1,
                    file_path=str(file_path),
                    error=str(e),
                    exc_info=True,
                )
            except ValueError as e:
                logger.error(
                    "value_error_adding_chunk",
                    chunk_index=i + 1,
                    file_path=str(file_path),
                    error=str(e),
                    exc_info=True,
                )
            except TypeError as e:
                logger.error(
                    "type_error_adding_chunk",
                    chunk_index=i + 1,
                    file_path=str(file_path),
                    error=str(e),
                    exc_info=True,
                )
            except Exception as _e:
                logger.critical(
                    "critical_unhandled_error_adding_chunk",
                    chunk_index=i + 1,
                    file_path=str(file_path),
                    exc_info=True,
                )
                # We don't re-raise here to allow processing of other chunks

        logger.info(
            "successfully_ingested_chunks",
            chunk_count=chunk_count,
            total_chunks=len(chunks),
            file_path=str(file_path),
        )
        return chunk_count > 0


async def load_existing_rag_documents(file_ingestor: IFileIngestor) -> int:
    """
    Load all existing documents from RAG directories into the knowledge graph.

    Args:
        file_ingestor: The file ingestor instance

    Returns:
        Number of documents processed
    """
    processed_count = 0

    # Process all knowledge base directories
    for knowledge_dir in settings.KNOWLEDGE_BASE_DIRS:
        knowledge_path = settings.BASE_DIR / knowledge_dir

        if not knowledge_path.exists():
            logger.warning("knowledge_base_directory_not_found", knowledge_path=str(knowledge_path))
            continue

        logger.info("processing_knowledge_base_directory", knowledge_path=str(knowledge_path))

        # Walk through all files in the directory tree
        for file_path in knowledge_path.rglob("*"):
            if file_path.is_file() and not file_path.name.startswith("."):
                # Check if file is in protected directories (should be processed)
                if is_path_in_knowledge_base(file_path):
                    try:
                        logger.info("loading_existing_document", filename=file_path.name)
                        await file_ingestor.ingest_file(file_path)
                        processed_count += 1
                    except FileProcessingError as e:
                        logger.error(
                            "failed_to_process_document",
                            file_path=str(file_path),
                            error=str(e),
                            exc_info=True,
                        )
                else:
                    logger.debug("skipping_protected_file", file_path=str(file_path))

    logger.info("loaded_existing_rag_documents", processed_count=processed_count)
    return processed_count


def create_file_ingestor(knowledge_graph: IKnowledgeGraph) -> FileIngestor:
    """
    Factory function to create a FileIngestor instance.

    Args:
        knowledge_graph: The knowledge graph service to use

    Returns:
        A configured FileIngestor instance
    """
    return FileIngestor(knowledge_graph)
